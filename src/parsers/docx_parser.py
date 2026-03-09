"""DOCX parser using python-docx."""
from docx import Document
from docx.shared import Inches
from typing import List, Dict, Any
import os

from src.parsers.base_parser import BaseParser, ParsedChunk


class DOCXParser(BaseParser):
    """Parser for DOCX documents."""
    
    def parse(self, file_path: str) -> List[ParsedChunk]:
        """
        Parse DOCX document and extract text with enhanced content handling.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of ParsedChunk objects
        """
        chunks = []
        
        try:
            doc = Document(file_path)
            
            # Extract content from different elements
            content_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append({
                        'type': 'paragraph',
                        'text': para.text.strip(),
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_content = []
                for row_idx, row in enumerate(table.rows):
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content.append(cell_text)
                    if row_content:
                        table_content.append(" | ".join(row_content))
                
                if table_content:
                    content_parts.append({
                        'type': 'table',
                        'text': "\n".join(table_content),
                        'table_number': table_idx + 1
                    })
            
            # Process content parts into chunks
            current_chunk_text = ""
            current_metadata = {
                'paragraph_count': len([p for p in content_parts if p['type'] == 'paragraph']),
                'table_count': len([p for p in content_parts if p['type'] == 'table']),
                'parser': 'python-docx',
                'total_parts': len(content_parts)
            }
            
            chunk_index = 0
            
            for part in content_parts:
                part_text = f"[{part['type'].upper()}] {part['text']}\n"
                
                # Check if adding this part would exceed chunk size
                if len(current_chunk_text + part_text) > 800 and current_chunk_text:
                    # Create chunk with current content
                    chunk = ParsedChunk(
                        text=self.clean_text(current_chunk_text),
                        page_number=None,  # DOCX doesn't have pages
                        chunk_index=chunk_index,
                        metadata=current_metadata.copy()
                    )
                    chunks.append(chunk)
                    
                    # Start new chunk
                    current_chunk_text = part_text
                    chunk_index += 1
                else:
                    current_chunk_text += part_text
            
            # Add remaining content as final chunk
            if current_chunk_text:
                chunk = ParsedChunk(
                    text=self.clean_text(current_chunk_text),
                    page_number=None,
                    chunk_index=chunk_index,
                    metadata=current_metadata
                )
                chunks.append(chunk)
        
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
        
        return chunks
    
    def extract_headings(self, file_path: str) -> List[str]:
        """
        Extract all headings from document.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of heading texts
        """
        try:
            doc = Document(file_path)
            headings = []
            
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    headings.append(para.text.strip())
            
            return headings
        except Exception:
            return []
